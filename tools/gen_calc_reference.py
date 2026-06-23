#!/usr/bin/env python3
"""
tools/gen_calc_reference.py — generate the component-calculation reference (.docx, native equations).
=====================================================================================================
Builds `docs/DCCREG_Turbine_component_calculations.docx`: for EVERY component in the netlist
(`topology_edge_list.csv` / the live `REF_MAP`), the governing formula(s) as NATIVE Word equations
(OMML, editable in Word — not images), the variable definitions, the value at the established anchor,
the dependency on the other components/solver fields, and the source (citation or in-repo derivation).

Pipeline: LaTeX -> MathML (latex2mathml) -> OMML (mathml2omml) -> embedded into python-docx. No Office
needed; the equations are real Word math (m:oMath) that open editable.

Firewall: pure EE/mechanical. This is a STUDY/REFERENCE artifact derived from the calculator sources
(index.html, sim/design_synth.py, reference/*.py); it changes no solver. Re-run to regenerate.
"""
import os

import latex2mathml.converter as _l2m
import mathml2omml
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
OUT = os.path.join(ROOT, "docs", "DCCREG_Turbine_component_calculations.docx")

_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUT = RGBColor(0x55, 0x60, 0x6B)
GOOD = RGBColor(0x1f, 0x5a, 0x2c)


def _omml(latex):
    """LaTeX -> list of native <m:oMath> elements."""
    mml = _l2m.convert(latex)
    omml = mathml2omml.convert(mml)
    xml = f'<root xmlns:m="{_M}" xmlns:w="{_W}">{omml}</root>'
    return list(etree.fromstring(xml))


# ---------------------------------------------------------------- doc helpers
doc = Document()
_norm = doc.styles["Normal"]
_norm.font.name = "Calibri"
_norm.font.size = Pt(10.5)


def eq(latex, center=True):
    """Display equation on its own line (native Word math)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    for el in _omml(latex):
        p._p.append(el)
    return p


def para(text="", *, bold=False, italic=False, size=None, color=None, space_after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
    return p


def runs(parts, space_after=6):
    """A paragraph built from (text, {opts}) runs — for inline mixed styling."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, opt in parts:
        r = p.add_run(text)
        r.bold = opt.get("bold", False); r.italic = opt.get("italic", False)
        if opt.get("size"): r.font.size = Pt(opt["size"])
        if opt.get("color"): r.font.color.rgb = opt["color"]
    return p


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def vartable(rows):
    """rows = [(symbol, meaning, value/anchor)]. A 3-col definitions table."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for c, txt in zip(hdr, ("symbol", "meaning", "value at the established anchor")):
        c.paragraphs[0].add_run(txt).bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
    for sym, mean, val in rows:
        cells = t.add_row().cells
        for c, txt in zip(cells, (sym, mean, val)):
            r = c.paragraphs[0].add_run(txt); r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def source(text):
    runs([("Source / derivation.  ", {"bold": True, "size": 9, "color": ACCENT}),
          (text, {"size": 9, "italic": True})], space_after=10)


def depends(text):
    runs([("In the calculator.  ", {"bold": True, "size": 9, "color": GOOD}),
          (text, {"size": 9})], space_after=4)


# ============================================================================
# TITLE
# ============================================================================
tp = para("DCCREG Turbine", bold=True, size=24, space_after=2)
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp = para("Component Calculation Reference", size=16, color=ACCENT, space_after=2)
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp2 = para("the governing formula, derivation, and dependency for every part in the netlist",
           italic=True, size=11, color=MUT, space_after=14)
sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER

para("Scope.", bold=True, space_after=2)
para("This document derives, from first principles or cited sources, the calculation for each of the 43 "
     "components in the machine netlist (topology_edge_list.csv) — the symmetric voltage-doubler core, the "
     "two resonant islands, the distributed-electromagnet reluctance motor, and the resonator tank. It is a "
     "paper-study companion to the calculator tools (index.html and the Pyodide synthesizer): every formula "
     "below is exactly the one the code evaluates, with the in-code symbol names. Equations are native Word "
     "math objects (editable), not images.")
para("Reading conventions.", bold=True, space_after=2)
para("Each component family gives: the governing equation(s); a table of variables with their value at the "
     "established design anchor (r_out 387 mm, V_target 15 kV, V_strike 20 kV, 3000 rpm); how the value is "
     "produced in the calculator (which live field feeds it, and what it depends on); and the source — a "
     "literature citation [n] (see References) or, for results derived in this project, the derivation plus the "
     "in-repo provenance. Tier tags follow the repo: [OC] derivable from standard physics, [IR] a design/"
     "empirical choice, [ME] a method.")
para("Firewall.", bold=True, space_after=2)
para("Pure electrical/mechanical engineering throughout — capacitor electrostatics, LC resonance, magnetics, "
     "spark-gap breakdown, field emission. No substrate or DCCREG theory enters any formula.")

# the netlist roster
h2("The netlist roster (43 components)")
para("The four subsystems and their reference designators (the names KiCad draws and the calculator binds):")
for fam, refs in [
    ("Doubler core", "C1, C2 (input varicaps) · Ca1, Cb1 (transfer caps) · SG1, SG2 (rail-return gaps)"),
    ("Resonant islands", "Cx3, Cx4 (pickup varicaps) · Lx3, Lx4 (resonant inductors) · "
                          "SG3a1, SG4a1 (load gaps) · SG3b1, SG4b1 (fire gaps) · BS3, BS4 (FE backstops)"),
    ("Distributed-electromagnet motor", "L_A1–6, L_B1–6 (stator coils) · "
                                         "C_AR1–6, C_BR1–6 (per-coil DC-block caps)"),
    ("Resonator tank", "L_R1, L_R2 (split coils) · C_R1 (tank cap)"),
]:
    runs([(f"{fam}:  ", {"bold": True, "size": 10}), (refs, {"size": 10})], space_after=3)

doc.add_page_break()

# ============================================================================
# PART 0 — FOUNDATIONS
# ============================================================================
h1("0  Foundations — the shared physics")
para("Three relations recur across every subsystem; they are stated once here and referenced below.")

h3("0.1  Physical constants")
vartable([
    ("ε₀", "vacuum permittivity", "8.8541878128×10⁻¹² F/m"),
    ("μ₀", "vacuum permeability", "4π×10⁻⁷ H/m"),
    ("ρ_Cu", "annealed-copper resistivity @20 °C", "1.68×10⁻⁸ Ω·m"),
])

h3("0.2  Parallel-plate capacitance")
para("Every fixed and variable capacitor in the machine is a parallel-plate (or sectored parallel-plate) "
     "structure: two conductive faces of area A separated by a dielectric of relative permittivity εr and "
     "thickness d.")
eq(r"C = \varepsilon_0\,\varepsilon_r\,\frac{A}{d}")
source("Standard electrostatics [1]. The fringing field is neglected (A ≫ d²), consistent with the brief's "
       "scope (fringing out of scope).")

h3("0.3  Charge-sharing (two-capacitor) loss")
para("When a charged capacitor is connected directly to another at a different voltage, the equalisation "
     "dissipates a fixed energy independent of the connecting resistance — the “two-capacitor paradox.” It is "
     "the tax the doubler pays on each bucket-brigade transfer, and the loss the resonant islands exist to "
     "avoid.")
eq(r"E_{loss} = \frac{1}{2}\,\frac{C_1 C_2}{C_1+C_2}\,(V_1-V_2)^2")
source("Classical result [2]; in-repo localisation in sim/brigade_tax_localize.py.")

h3("0.4  LC resonance, characteristic impedance, quality factor")
para("The islands, the motor coils, and the tank are all LC circuits. For inductance L and capacitance C in "
     "series with loss R:")
eq(r"\omega_0 = \frac{1}{\sqrt{LC}}, \quad f_0=\frac{\omega_0}{2\pi}, \quad "
   r"Z_0=\sqrt{\frac{L}{C}}, \quad Q=\frac{Z_0}{R}=\frac{\omega_0 L}{R}")
source("Standard linear circuits [2].")

doc.add_page_break()

# ============================================================================
# PART 1 — DOUBLER CORE
# ============================================================================
h1("1  Doubler core — C1, C2, Ca1, Cb1, SG1, SG2")
para("The heart of the machine is a symmetric (Bennet) voltage doubler: two rotating variable capacitors "
     "(C1, C2) are charged at high capacitance and discharged at low capacitance, the rail-return gaps "
     "(SG1, SG2) rectify the swing, and the transfer caps (Ca1, Cb1) carry charge between the two halves. "
     "Mechanical work modulating C against the trapped charge is pumped up in voltage each cycle.")

h2("1.1  C1, C2 — the input variable capacitors (sectored varicaps)")
runs([("Role.  ", {"bold": True}),
      ("Rotor/stator sectored plates whose overlap — and thus capacitance — varies between C_min "
       "(anti-aligned) and C_max (aligned) as the rotor turns. The modulation is the pump's drive.",
       {})], space_after=4)
para("The aligned (maximum) capacitance is the sectored parallel-plate area over the vertical gap. With "
     "N_sec sectors of which n_kept are active over the active band r_in…r_out:")
eq(r"A_{kept} = n_{kept}\,\frac{1}{N_{sec}}\,\pi\,\big(r_{out}^2 - r_{in}^2\big)")
eq(r"C_{max} = \varepsilon_0\,\frac{A_{kept}}{g_v}")
vartable([
    ("r_in, r_out", "active-band inner / outer radius", "95 mm / 387 mm"),
    ("N_sec, n_kept", "total sectors / active sectors", "12 / 6"),
    ("g_v", "vertical plate gap", "7 mm"),
    ("C_max", "aligned capacitance (computed)", "≈ 280 pF"),
    ("C_min", "anti-aligned capacitance (incl. stray floor)", "16 pF"),
])
depends("Cmax_from_geom(r_in, r_out, g_v, n_kept, N_sec) in sim/design_synth.py; the schematic slot sv_C1 / "
        "sv_C2 shows C1max_pF live. C1/C2 feed the pump-gain solver (§1.3).")
source("Sectored parallel-plate area + §0.2 [1]; the sectored-disc construction follows "
       "reference/SectoredDiscCalculator.jsx and docs/brief-blockC1-geometry-to-rotorcap.md [OC].")

h2("1.2  Ca1, Cb1 — the bucket-brigade transfer capacitors")
runs([("Role.  ", {"bold": True}),
      ("A solid annular bus ring on the back of each stator plate forms a fixed capacitor (Mylar dielectric) "
       "that carries charge between the doubler halves each stroke.", {})], space_after=4)
eq(r"C_a = \varepsilon_0\,\varepsilon_r\,\frac{A_{ring}}{t_{Mylar}}")
vartable([
    ("A_ring", "annular bus-ring electrode area", "from active-band radii"),
    ("ε_r, t", "Mylar permittivity / thickness", "≈ 3.1 / design"),
    ("C_a = C_b", "transfer capacitance", "309 pF"),
])
runs([("The transfer tax.  ", {"bold": True}),
      ("Each Ca/Cb hand-off is a charge-sharing equalisation, so it pays the §0.3 loss; summed over the two "
       "dominant phase transfers this is 9.79 mJ/fire (≈ 69 % of the cycle tax). This is intrinsic to the "
       "ratchet — see §1.3.", {})], space_after=4)
depends("transferCaps(s) in index.html (C = ε0·εr·A/t, solid annulus); sv_Ca1 / sv_Cb1 show Ca_pF live; "
        "Ca/Cb set the doubler matrix in §1.3 and the brigade tax in §0.3.")
source("§0.2 parallel-plate [1]; the two-cap tax §0.3 [2]; per-transfer localisation in "
       "sim/brigade_tax_localize.py (Σ = 19.58 mJ/cycle) [OC].")

h2("1.3  Pump gain z and net-electrical fraction η")
para("The doubler is solved as a charge-defined capacitor network: from the node voltages V = (v1,v2,v3,v4) "
     "the trapped charges are")
eq(r"Q_1=(C_1+C_{par})v_1+C_a(v_1-v_2), \;\; Q_4=(C_2+C_{par})v_4+C_b(v_4-v_3),\;\dots")
para("each stroke re-solves the network at the new capacitances subject to the rail-return rectification, and "
     "the steady per-cycle voltage growth is the pump gain z; the fraction of the mechanical input that leaves "
     "as net electrical energy is η:")
eq(r"z = \frac{|v_1|+|v_4|}{|v_1'|+|v_4'|}\Big|_{\text{steady}}, \qquad "
   r"\eta = \operatorname{median}\!\left(\frac{\Delta U}{W_{mech}}\right)")
vartable([
    ("z", "per-cycle pump gain (voltage ratio)", "1.334"),
    ("η", "net-electrical fraction at the 4-node point", "0.386"),
    ("C_par", "parasitic/stray capacitance floor", "20 pF"),
])
runs([("Why these values, not larger.  ", {"bold": True}),
      ("The same rail-return rectification that makes the ladder pump (conduction at the current zero each "
       "cycle) is exactly what forbids recovering the transfer tax by over-resonating the core — “the "
       "equalisation IS the pump.” So the core fraction is fixed near η = 0.386; recovery is only available on "
       "downstream sinks (the islands, §2.2). The machine's operating efficiency is therefore", {}),
      ], space_after=2)
eq(r"\eta_{op}=\frac{E_{useful}+f_{rec}\,E_{island\text{-}tax}}{E_{useful}+E_{doubler\text{-}tax}+E_{island\text{-}tax}}"
   r"=\frac{6.153+f_{rec}\cdot4.407}{20.347}\approx 0.52")
depends("solve_doubler4 / charges_from_voltages / solve_phase in reference/doubler_core.py (the frozen "
        "authority, mirrored from index.html); z_eta_Wmech and operating_point in sim/design_synth.py.")
source("Symmetric electrostatic voltage doubler — Bennet's doubler [4], analysed in the charge-pump framework "
       "of de Queiroz [3]; z is reproduced to 0.002 % by the de Queiroz analytic eigen-matrix "
       "(xsim_queiroz_matrix.py) and to a few % by an independent ngspice charge-defined varicap. The "
       "efficiency composition is derived and recorded in docs/efficiency-resolution.md [OC/ME].")

h2("1.4  SG1, SG2 — the rail-return spark gaps")
para("The doubler has no diodes; rectification is by self-breaking spark gaps that hold off until the gap "
     "voltage reaches the strike level, then arc and self-quench. The strike voltage is the breakdown of the "
     "gap (sphere gap in vacuum/low-pressure cavity):")
eq(r"V_{strike} = E_{bd}\,g_{sg} \quad\text{(field form)}, \qquad V_{bd}=K\,g_{sg}^{\,0.6}\ \text{(vacuum law)}")
vartable([
    ("g_sg", "spark-gap spacing", "0.5 mm (rail gaps)"),
    ("E_bd", "breakdown field (W-Cu spheres)", "≈ 3 kV/mm"),
    ("K", "vacuum total-voltage coefficient", "60 (kV·mm⁻⁰·⁶)"),
    ("V_strike", "established strike voltage", "20 kV"),
])
depends("Gaps carry V_strike_kV live (sv_SG1 / sv_SG2); the holdoff threshold sets the doubler rectification "
        "in §1.3 and the firing reach in §2.4. V_bd guards insulation invariant I4 in design_synth.py.")
source("High-voltage spark-gap breakdown [6]; the vacuum total-voltage law V_bd ∝ g^0.6 is the Cranberg/area "
       "form [6], used as design rule K·g^0.6 [IR]. Sphere field 3 kV/mm from the freeze (doc §5) [IR].")

doc.add_page_break()

# ============================================================================
# PART 2 — RESONANT ISLANDS
# ============================================================================
h1("2  Resonant islands — Cx3/Cx4, Lx3/Lx4, SG3a1/SG4a1, SG3b1/SG4b1, BS3/BS4")
para("Each island is a flying-bucket pickup capacitor (Cx) that is charged from the rail through a load gap "
     "(SGa), rings its charge into the receiving bank through a series inductor (Lx) in a half-cycle resonant "
     "transfer, fires through a fire gap (SGb) at strike, and is bounded by a field-emission backstop (BS). "
     "The series Lx is the efficiency fix: it turns the lossy direct dump (§0.3) into a near-lossless resonant "
     "swap.")

h2("2.1  Cx3, Cx4 — the island pickup varicaps")
runs([("Role.  ", {"bold": True}),
      ("Rotating island plates that pick up charge from the rail at high capacitance; the same sectored "
       "parallel-plate geometry as C1/C2 (§1.1), sized to the island band.", {})], space_after=4)
eq(r"C_{x,max} = \varepsilon_0\,\frac{A_{island}}{g_{isl}}")
vartable([
    ("A_island", "island plate overlap area", "from island band"),
    ("g_isl", "island plate gap", "4 mm"),
    ("C_x,max", "island pickup capacitance", "471 pF"),
])
depends("Cx_maxMm field feeds sv_Cx3 / sv_Cx4 and the LC ring (§2.2) and the shuttle (§2.4).")
source("§0.2 / §1.1 sectored parallel-plate [1] [OC].")

h2("2.2  Lx3, Lx4 — the series resonant inductors (the half-cycle transfer)")
para("With Lx in series, the island (C_src) discharges into the bank (C_bank) as an underdamped series-RLC "
     "ring. The transfer completes in one half-cycle, at the first current zero, where the inductor hands all "
     "its energy to the bank. The effective series capacitance, ring impedance, half-period and peak current:")
eq(r"C_{eff}=\frac{C_{src}\,C_{bank}}{C_{src}+C_{bank}}")
eq(r"Z=\sqrt{\frac{L_x}{C_{eff}}}, \qquad t_{1/2}=\pi\sqrt{L_x C_{eff}}, \qquad "
   r"i_{pk}=\frac{\Delta V}{Z}=\Delta V\sqrt{\frac{C_{eff}}{L_x}}")
para("The energy a direct dump would lose (§0.3) is E_2cap; the resonant ring instead loses only the resistive "
     "dissipation over the half-cycle, which vanishes as the ring Q rises — recovering a fraction f_rec of the "
     "tax:")
eq(r"E_{2cap}=\tfrac{1}{2}C_{eff}\,\Delta V^2, \qquad "
   r"E_{loss}=\frac{\pi}{2Q}\,E_{2cap}, \qquad f_{rec}=1-\frac{\pi}{2Q}")
vartable([
    ("L_x", "series island inductor", "1 mH"),
    ("C_src / C_bank", "island / receiving bank", "471 pF / 2.64 nF"),
    ("ΔV", "island-to-bank voltage step", "≈ 5 kV"),
    ("Q = Z/R", "ring quality factor", "≈ 729 (R = 2 Ω)"),
    ("t_1/2", "half-cycle transfer time", "≈ 0.8 µs"),
    ("f_rec", "recovered-tax fraction", "≈ 0.996"),
])
depends("closed_form(...) and the RK4 integrate(...) in reference/island_resonant_core.py; the live tool calls "
        "irc.integrate to get f_rec, which sets η_op (§1.3). sv_Lx3 / sv_Lx4 show Lx_mH. t_1/2 must fit the "
        "spark-gap conduction window (invariant I10).")
source("Standard series-RLC resonant charge transfer [2]; the loss→0 as Q→∞ and f_rec are derived in "
       "reference/island_resonant_core.py and cross-checked by an independent ngspice LC ring to <0.6 % "
       "(NGSPICE-CONFIRMS, S2). The resonant-transfer principle follows de Queiroz [3] [OC].")

h2("2.3  SG3a1/SG4a1 (load) and SG3b1/SG4b1 (fire) — the island gaps")
para("Four self-breaking spark gaps clock the island: the load gaps SGa admit charge from the rail to the "
     "island; the fire gaps SGb release the island into the cross-coupled bank when the collapsing-capacitance "
     "voltage reaches strike. All obey the breakdown of §1.4:")
eq(r"V_{strike} = E_{bd}\,g_{sg}")
vartable([
    ("g_sg", "island gap (sphere)", "ball d = 12 mm, freeze §5"),
    ("V_strike", "strike voltage", "20 kV"),
    ("V_ceil", "insulation ceiling", "21 kV"),
])
depends("All four carry V_strike_kV live; the fire gaps gate the shuttle reach (§2.4). Gap placement radius "
        "r_gap and ball d set the electrode-overlap timing window (firing-geometry, invariants I11/I12).")
source("Spark-gap breakdown [6]; gap roles documented in docs/kicad/gap-topology-of-record.md [IR/OC].")

h2("2.4  BS3, BS4 — the Fowler–Nordheim field-emission backstops")
para("In parallel with the fire gaps, a field-emission backstop bleeds charge softly once the field exceeds an "
     "onset (here 0.6·V_strike). The emission current follows the Fowler–Nordheim law (a cold-cathode "
     "tunnelling current), and the bled energy is the time integral of V·I over the swing and dwell:")
eq(r"I_{FN}(V)=A\,V^2\,\exp\!\left(-\frac{B}{V}\right), \qquad B=k\,V_{strike}")
eq(r"E_{FE}=\int V\,I_{FN}(V)\,dt")
vartable([
    ("onset", "FE turn-on voltage", "0.6·V_strike = 12 kV"),
    ("I_ref", "designed leakage at V_strike", "30 µA"),
    ("A, B", "FN coefficients (B = k·V_strike)", "fit to I_ref, k = 3"),
])
depends("fn_coeffs / fn_current / fe_arc_budget in reference/commutator_real_core.py; sv_BS3 / sv_BS4 show "
        "0.6·V_strike. E_FE is a real, perturbable loss (it is the bounded price of holding the gap off).")
source("Fowler & Nordheim, cold field electron emission [5]; the budget integral and arc accounting are "
       "in reference/commutator_real_core.py [OC for the FN law; IR for the coefficient fit].")

h2("2.5  Shuttle reach — V*, E_fire, C_fire (how the island fires)")
para("When the finite rail couples its charge onto the larger island it shares to a diluted common voltage V*; "
     "the mechanical capacitance collapse then raises that voltage toward strike, where the fire energy and "
     "fire-state capacitance are set by the charge delivered:")
eq(r"V^{*}=\frac{C_{rail}\,V_{rail}}{C_{rail}+C_x}, \qquad "
   r"E_{fire}=\tfrac{1}{2}\,Q\,V_{strike}, \qquad C_{fire}=\frac{Q}{V_{strike}}")
depends("shuttle_Wcoll(Cx) wraps the frozen shuttle (island_charging_cosim); design_synth invariant I10 "
        "requires V_strike > V* (the collapse must reach strike) and the ring t_1/2 to fit the SG window.")
source("Charge-conserving capacitor coupling + collapse [1] [2]; iteration in sim/island_charging_cosim.py "
       "[OC].")

doc.add_page_break()

# ============================================================================
# PART 3 — MOTOR
# ============================================================================
h1("3  Distributed-electromagnet motor — L_A1–6, L_B1–6, C_AR1–6, C_BR1–6")
para("Twelve stator electromagnet coils (two push-pull groups of six, L_A and L_B) spin the rotor as a "
     "reluctance motor. Each coil is series-resonated by its own per-coil DC-block capacitor (C_AR/C_BR), which "
     "also blocks the DC rail (a coil is a near short at DC). The drive frequency is set by the firing rate.")

h2("3.1  Drive frequency")
eq(r"PRF=\frac{N_{ev}\cdot n_{rpm}}{60}, \qquad f_{drive}=\tfrac{1}{2}\,PRF\ \ (\text{push-pull, alternate strokes})")
vartable([
    ("N_ev", "commutation events per revolution", "design"),
    ("n_rpm", "rotor speed", "3000 rpm"),
    ("f_drive", "coil drive frequency", "≈ 150–300 Hz"),
])
source("Kinematic definition of pulse-repetition frequency; demMotor(s) in index.html [OC].")

h2("3.2  L_A1–6, L_B1–6 — the stator coils (reluctance inductance)")
para("Each coil is wound on a C-core with an air gap facing the rotor pole. Its inductance is the "
     "magnetic-circuit (reluctance) result for N turns over a pole-face area A_gap and total air gap l_gap, and "
     "it is sized so the coil resonates with its block cap at f_drive:")
eq(r"L=\frac{\mu_0\,N^2\,A_{gap}}{l_{gap}} \;\Longrightarrow\; N=\sqrt{\frac{L\,l_{gap}}{\mu_0\,A_{gap}}}, "
   r"\qquad L=\frac{1}{\omega^2 C_{block}}")
para("The coil resistance (copper) and its loaded quality factor:")
eq(r"R_{coil}=\rho_{Cu}\,\frac{l_{wire}}{A_{cond}}, \qquad Z_0=\sqrt{\frac{L}{C_{block}}}, \qquad Q=\frac{Z_0}{R_{coil}}")
vartable([
    ("A_gap", "pole-face area", "design (cm²)"),
    ("l_gap", "total air gap (n_gaps × clearance)", "design (mm)"),
    ("L", "coil inductance (resonant with C_block)", "≈ 0.64 H"),
    ("N", "turns", "from L, l_gap, A_gap"),
])
depends("demMotor(s) in index.html computes L = 1/(ω²C), N = √(L·l_gap/(μ0·A_gap)), Z0, R_coil, Q. "
        "sv_L_A* / sv_L_B* show the design inductance (Block-D static — not computed by the live design_synth).")
source("Magnetic-circuit / reluctance inductance L = μ0 N² A/l_gap [7]; series resonance §0.4 [2]; "
       "conductor resistance and AWG (§3.3, §A) [OC].")

h2("3.3  C_AR1–6, C_BR1–6 — the per-coil DC-block / resonating capacitors")
para("Each coil's series capacitor performs two jobs: it blocks the DC rail (so the near-short coil does not "
     "crowbar it) and it resonates the coil at the drive frequency. Its value is the series-resonance condition, "
     "and the AC ampere-turns the branch can drive — the torque-bearing quantity — follow a closed form that is "
     "independent of frequency and turns:")
eq(r"C_{block}=\frac{1}{\omega^2 L}=\frac{1}{(2\pi f_{drive})^2\,L}")
eq(r"N I=(V_{rating}-V_{bias})\sqrt{\frac{C_{block}\,l_{gap}}{\mu_0\,A_{gap}}}, \qquad "
   r"V_{cap,pk}=V_{bias}+Q\,V_{ripple}")
vartable([
    ("C_block", "per-coil series cap", "440 nF (target)"),
    ("V_rating", "cap voltage rating", "20 kV"),
    ("V_bias", "DC bias across the cap", "design"),
    ("N·I", "AC ampere-turns (∝ torque per branch)", "closed form"),
])
runs([("Energy / safety.  ", {"bold": True}),
      ("Each cap stores ½·C·V², a lethal kJ-class store on the spinning stator — which drives the per-coil "
       "vs per-group topology choice (same total energy either way).", {})], space_after=4)
depends("demMotor(s) §4 in index.html: the cap rating is the binding limit; sv_C_AR* / sv_C_BR* show the "
        "440 nF design value (Block-D static).")
source("Series resonance §0.4 [2]; the ampere-turn closed form N·I = (V_rating−V_bias)·√(C·l_gap/(μ0·A_gap)) "
       "is derived in index.html demMotor / the Block-D brief from L = μ0 N²A/l_gap and Z0 = √(L/C) [OC/ME].")

doc.add_page_break()

# ============================================================================
# PART 4 — RESONATOR
# ============================================================================
h1("4  Resonator tank — L_R1, L_R2, C_R1")
para("The output stage is an LC tank: two conical coils in series-aiding (L_R1, L_R2) resonate with a "
     "through-mica inter-electrode capacitor (C_R1), pulse-excited by the firing rate. It is built from "
     "circular-loop inductance, parallel-plate capacitance, and the LC relations of §0.4.")

h2("4.1  C_R1 — the through-septum tank capacitor")
eq(r"C_R=\varepsilon_0\,\varepsilon_d\,\frac{A_{align}}{h_{disc}}")
vartable([
    ("A_align", "aligned rotor-face area through the disc", "full rotor face"),
    ("ε_d, h_disc", "mica/garolite permittivity / disc thickness", "design"),
    ("C_R", "tank capacitance", "789 pF"),
])
depends("resonatorCore(s) in index.html: C_R = ε0·εd·A_align/h_disc; sv_C_R1 shows C_R_pF live. "
        "C_R sets f0 and Z0 with L (below).")
source("§0.2 parallel-plate through the septum dielectric [1] [OC].")

h2("4.2  L_R1, L_R2 — the split conical coils (loop inductance)")
para("The tank inductance is summed over the winding's circular loops: each loop contributes its high-"
     "frequency self-inductance, and every pair contributes Maxwell's mutual inductance of two coaxial loops "
     "(via the complete elliptic integrals K, E). Series-aiding, the two cones add with their mutual.")
eq(r"L_{self}=\mu_0\,r\left(\ln\frac{8r}{a}-2\right)")
eq(r"k^2=\frac{4\,r_1 r_2}{(r_1+r_2)^2+d^2}, \qquad "
   r"M=\mu_0\sqrt{r_1 r_2}\left[\left(\tfrac{2}{k}-k\right)K(k)-\tfrac{2}{k}\,E(k)\right]")
eq(r"L=\sum_i L_{self,i}+2\!\!\sum_{i<j}\! M_{ij}")
vartable([
    ("r, a", "loop radius / wire radius", "from cone geometry"),
    ("d", "axial spacing between two loops", "winding pitch"),
    ("L", "total tank inductance (both cones)", "≈ 79 µH (≈ 39.5 µH each)"),
])
depends("Lself(r,a), Mloop(r1,r2,d) via ellipKE, and L = Σself + 2ΣΣM in index.html resonatorCore; "
        "sv_L_R1 / sv_L_R2 show the per-coil design inductance (≈ 39.5 µH).")
source("HF self-inductance of a circular loop, L = μ0 r(ln(8r/a) − 2) [8]; Maxwell's mutual inductance of "
       "coaxial circular loops via complete elliptic integrals K, E [8] [OC].")

h2("4.3  Tank resonance, impedance, Q, and the damped drive")
para("With the total L and C_tot = C_R + C_self + stray, the tank resonance, characteristic impedance and "
     "copper-limited Q (through the skin-effect AC resistance) are:")
eq(r"f_0=\frac{1}{2\pi\sqrt{L\,C_{tot}}}, \qquad Z_0=\sqrt{\frac{L}{C_{tot}}}, \qquad Q=\frac{\omega_0 L}{R_{ac}}")
eq(r"\delta=\sqrt{\frac{\rho_{Cu}}{\pi f_0 \mu_0}}, \qquad R_{ac}=\frac{\rho_{Cu}\,l_{wire}}{\pi\,d_{OD}\,\delta}, "
   r"\qquad f_d=f_0\sqrt{1-\frac{1}{4Q^2}}")
vartable([
    ("C_tot", "tank C incl. self-capacitance + stray", "≈ 789 pF + …"),
    ("f_0", "tank resonant frequency", "637 kHz"),
    ("δ", "copper skin depth at f0", "≈ 80 µm"),
    ("f_d", "damped natural frequency", "≈ f0"),
])
para("The single-layer self-capacitance of the coil uses the Medhurst estimate (D, H in cm):")
eq(r"C_{self}=D\left(0.1126\,\frac{H}{D}+0.08+\frac{0.27}{\sqrt{H/D}}\right)")
depends("resonatorCore(s) in index.html: f0, Z0, Q, skin depth δ, Rac, fd, and the Medhurst Cself; the badge "
        "f0/Z0 are surfaced live in Block-R.")
source("LC resonance and skin effect §0.4 [2]; Medhurst single-layer self-capacitance [9] [OC/IR].")

doc.add_page_break()

# ============================================================================
# PART 5 — DEPENDENCY MAP
# ============================================================================
h1("5  How the components depend on each other (the calculator chain)")
para("The calculators are producer/consumer: geometry produces capacitances, the frozen doubler consumes them "
     "to produce z/η, the islands recover part of the tax, and a battery of invariants couples the whole. The "
     "central dependencies:")
for a, b in [
    ("Geometry (r_in, r_out, g_v, N_sec)", "→ C1/C2 C_max (§1.1) and Cx (§2.1)"),
    ("C1, C2, Ca, Cb, C_par", "→ pump gain z and η via the doubler network (§1.3)"),
    ("Cx, L_x, C_bank, ΔV, R", "→ resonant transfer t_1/2, i_pk, f_rec (§2.2) → η_op (§1.3)"),
    ("V_strike, g_sg, ball d", "→ rail/island rectification (§1.4, §2.3) and shuttle reach V* (§2.5)"),
    ("f_drive, C_block", "→ coil L (§3.2) → turns N, Z0, Q, N·I (§3.2–3.3)"),
    ("L (loops), C_R, C_self", "→ tank f0, Z0, Q (§4.3)"),
]:
    runs([(a + "  ", {"bold": True, "size": 9.5, "color": ACCENT}), (b, {"size": 9.5})], space_after=3)

para("")
para("The invariants that bind them (design_synth.py I1–I13).", bold=True, space_after=2)
para("I2 frozen-solver authority (z, η, W_coll come only from the solvers) · I3 z within the validated band "
     "[1.20, 1.45] · I4 insulate-first (gap V_bd > V_target, septum holds) · I5 tax managed (η ≥ 0.15) · "
     "I6 parasitic floor (C_par ≥ 20 pF) · I7 motor matched (output ≤ pump net; f_res = PRF) · I9 mechanical "
     "(rim speed < 200 m/s) · I10 shuttle integrity (V_strike < ceiling, V* reaches strike, ring t_1/2 fits the "
     "SG window) · I11/I12 firing geometry (electrode-overlap window, cross-fire margin) · I13 island recovery. "
     "A design is feasible only inside the region all of these allow; the binding one is what limits it.")
para("Producer/consumer discipline.", bold=True, space_after=2)
para("The geometry/motor/resonator blocks are producers: they compute capacitances/inductances and feed them "
     "in, but never write the rotor-cap fields and never call the doubler solver. The frozen doubler "
     "(reference/doubler_core.py) is the sole consumer/authority on z and η. This is what keeps the flexible "
     "dimensions from quietly overriding a derivation.")

doc.add_page_break()

# ============================================================================
# REFERENCES
# ============================================================================
h1("References")
refs = [
    "[1] D. J. Griffiths, Introduction to Electrodynamics, 4th ed., Cambridge Univ. Press, 2017 — "
    "parallel-plate capacitance, charge-conserving coupling.",
    "[2] C. K. Alexander & M. N. O. Sadiku, Fundamentals of Electric Circuits, 6th ed., McGraw-Hill, 2017 — "
    "LC resonance, characteristic impedance, quality factor, skin effect, the two-capacitor charge-sharing loss.",
    "[3] A. C. M. de Queiroz, “Electrostatic vibration energy harvesters with separated charge injectors / "
    "generalized symmetric charge-pump multipliers,” IEEE — the charge-defined varicap doubler analysis used as "
    "the analytic witness for z (eigen-matrix; xsim_queiroz_matrix.py).",
    "[4] Abraham Bennet, “Description of a new electrical doubler,” Phil. Trans. R. Soc. 77, 288 (1787) — the "
    "symmetric electrostatic voltage doubler.",
    "[5] R. H. Fowler & L. W. Nordheim, “Electron emission in intense electric fields,” Proc. R. Soc. Lond. A "
    "119, 173–181 (1928) — the cold field-emission (tunnelling) current law I ∝ V² exp(−B/V).",
    "[6] M. S. Naidu & V. Kamaraju, High Voltage Engineering, 5th ed., McGraw-Hill, 2013 — spark-gap and "
    "sphere-gap breakdown; the vacuum total-voltage breakdown law V_bd ∝ d^0.6 (Cranberg/area effect).",
    "[7] A. E. Fitzgerald, C. Kingsley, S. D. Umans, Electric Machinery, 6th ed., McGraw-Hill, 2003 — "
    "magnetic-circuit / reluctance inductance L = μ0 N² A / l_gap.",
    "[8] F. W. Grover, Inductance Calculations: Working Formulas and Tables, Dover, 1946 — Maxwell's mutual "
    "inductance of coaxial circular loops (complete elliptic integrals) and the high-frequency self-inductance "
    "of a circular loop.",
    "[9] R. G. Medhurst, “H.F. resistance and self-capacitance of single-layer solenoids,” Wireless Engineer "
    "24, 35–43 & 80–92 (1947) — the single-layer self-capacitance estimate.",
    "[R] In-repo derivations and validations (provenance for the project-specific composite results): "
    "reference/doubler_core.py, reference/island_resonant_core.py, reference/commutator_real_core.py, "
    "sim/design_synth.py, index.html (demMotor / resonatorCore / transferCaps), docs/efficiency-resolution.md, "
    "and the CHANGELOG verdict trail (NGSPICE-CONFIRMS, RESONANT-TRANSFER-MODELED, EFFICIENCY-RESOLVED).",
]
for r in refs:
    para(r, size=9, space_after=5)

para("")
runs([("Generated by ", {"size": 8, "color": MUT, "italic": True}),
      ("tools/gen_calc_reference.py", {"size": 8, "color": MUT}),
      (" — re-run to regenerate from the calculator sources. Equations are native Word math (OMML).",
       {"size": 8, "color": MUT, "italic": True})], space_after=2)

doc.save(OUT)
print("wrote", OUT)
# report equation count
import zipfile
xml = zipfile.ZipFile(OUT).read("word/document.xml").decode()
print("native OMML equations:", xml.count("<m:oMath"))
print("headings:", xml.count('<w:pStyle w:val="Heading'))
